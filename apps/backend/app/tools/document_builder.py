"""Resume document generation (PDF + DOCX) and GCS upload."""
from __future__ import annotations

import html as _html
import logging
import re
from datetime import date, timedelta

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from app.core.config import get_settings
from app.core.gcp import get_gcs_client
from app.schemas.rendered_resume import ContactLink, StructuredResume
from app.tools.resume_render import (
    DEFAULT_TEMPLATE,
    build_structured_resume,
    html_to_pdf,
    render_resume_pdf,
)

logger = logging.getLogger(__name__)

SIGNED_URL_EXPIRY = timedelta(days=7)


async def build_resume_pdf(
    resume_text: str,
    output_path: str,
    *,
    template_id: str = DEFAULT_TEMPLATE,
    structured: StructuredResume | None = None,
    candidate_name: str = "Candidate",
) -> str:
    """Render the tailored resume to a styled PDF using one of the HTML/CSS
    templates (Chromium via Playwright).

    Prefers a `structured` resume (from the resume_structurer agent). When it's
    not supplied, falls back to deterministically parsing the writer's
    plain-text output. Async because the renderer drives a headless browser.
    """
    if structured is None:
        structured = build_structured_resume(resume_text, candidate_name)
    return await render_resume_pdf(structured, template_id, output_path)


def _add_bottom_border(paragraph, color: str = "16213e") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


_CL_ACCENT = "#1A5C96"
_CL_GREY = "#555555"
_CL_DATE_GREY = "#444444"
_SIGNOFF_RE = re.compile(
    r"^(sincerely|regards|best regards|best|warm regards|kind regards|"
    r"yours (truly|sincerely)|respectfully|thank you)[,.]?$",
    re.IGNORECASE,
)


def _clean_cover_body(text: str, candidate_name: str) -> list[str]:
    """Split the agent's letter into body paragraphs, dropping any salutation
    or sign-off it included (the template renders those itself for a
    consistent look)."""
    paras = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    # Drop a leading "Dear ..." salutation.
    if paras and paras[0].lower().startswith("dear "):
        paras = paras[1:]
    # Drop trailing sign-off lines ("Sincerely," / the candidate's name / a
    # short title line) from the end.
    name_low = candidate_name.strip().lower()
    while paras:
        last = paras[-1].strip()
        low = last.lower()
        is_signoff = bool(_SIGNOFF_RE.match(low))
        is_name = low == name_low
        is_short_tag = len(last.split()) <= 5 and "." not in last and "," in last
        if is_signoff or is_name or (is_short_tag and len(paras) > 1):
            paras.pop()
        else:
            break
    return paras


def _cl_contact_line(links: list[ContactLink], location: str | None) -> str:
    parts: list[str] = []
    for ln in links:
        if ln.kind in ("email", "phone"):
            parts.append(_html.escape(ln.label))
    if location:
        parts.append(_html.escape(location))
    for ln in links:
        if ln.kind not in ("email", "phone"):
            parts.append(_html.escape(ln.label))
    return "  |  ".join(parts)


async def build_cover_letter_pdf(
    cover_letter_text: str,
    candidate_name: str,
    role_title: str,
    company_name: str,
    output_path: str,
    *,
    links: list[ContactLink] | None = None,
    location: str | None = None,
    headline: str | None = None,
) -> str:
    """Render a clean, professional cover-letter PDF (Calibri, blue accent,
    contact rule, recipient block, justified body) via headless Chromium."""
    body_paras = _clean_cover_body(cover_letter_text, candidate_name)
    body_html = "".join(f"<p>{_html.escape(p)}</p>" for p in body_paras)
    contact = _cl_contact_line(links or [], location)
    today = date.today().strftime("%B %d, %Y")
    tagline = headline or role_title

    page_html = f"""<!doctype html><html><head><meta charset="utf-8"><style>
@page {{ size: A4; margin: 18mm 18mm; }}
* {{ box-sizing: border-box; }}
html, body {{ margin:0; padding:0; }}
body {{ font-family:'Calibri','Segoe UI',Helvetica,Arial,sans-serif; font-size:11pt;
       line-height:1.45; color:#222; }}
.name {{ font-size:18pt; font-weight:bold; color:{_CL_ACCENT}; }}
.contact {{ font-size:10pt; color:{_CL_GREY}; padding-bottom:5px;
           border-bottom:1px solid {_CL_ACCENT}; margin-bottom:16px; }}
.date {{ color:{_CL_DATE_GREY}; margin-bottom:14px; }}
.recipient {{ margin-bottom:14px; }}
.recipient .who {{ font-weight:bold; }}
.subject {{ font-size:12pt; font-weight:bold; color:{_CL_ACCENT}; margin-bottom:12px; }}
.salutation {{ margin-bottom:12px; }}
.body p {{ text-align:justify; margin:0 0 11px; }}
.closing {{ margin-top:14px; }}
.sign-name {{ font-weight:bold; color:{_CL_ACCENT}; }}
.sign-tag {{ font-size:10pt; color:{_CL_GREY}; }}
</style></head><body>
<div class="name">{_html.escape(candidate_name)}</div>
<div class="contact">{contact}</div>
<div class="date">{today}</div>
<div class="recipient">
  <div class="who">Hiring Manager</div>
  <div>{_html.escape(company_name)}</div>
</div>
<div class="subject">Re: Application for {_html.escape(role_title)}</div>
<div class="salutation">Dear Hiring Manager,</div>
<div class="body">{body_html}</div>
<div class="closing">Warm regards,</div>
<div class="sign-name">{_html.escape(candidate_name)}</div>
{f'<div class="sign-tag">{_html.escape(tagline)}</div>' if tagline else ''}
</body></html>"""

    return await html_to_pdf(page_html, output_path)


def _cl_contact_text(links: list[ContactLink], location: str | None) -> str:
    """Plain-text contact line for the DOCX cover letter (email/phone first,
    then location, then profile links) — mirrors the PDF's ordering."""
    parts: list[str] = []
    for ln in links:
        if ln.kind in ("email", "phone"):
            parts.append(ln.label)
    if location:
        parts.append(location)
    for ln in links:
        if ln.kind not in ("email", "phone"):
            parts.append(ln.label)
    return "  |  ".join(parts)


def build_cover_letter_docx(
    cover_letter_text: str,
    candidate_name: str,
    role_title: str,
    company_name: str,
    output_path: str,
    *,
    links: list[ContactLink] | None = None,
    location: str | None = None,
    headline: str | None = None,
) -> str:
    """Edit-ready Word version of the cover letter, styled to match the PDF
    (Calibri, blue accent name/subject/signature, contact rule, recipient
    block, justified body). Synchronous python-docx — run in a threadpool from
    async callers if needed."""
    body_paras = _clean_cover_body(cover_letter_text, candidate_name)
    accent = RGBColor(0x1A, 0x5C, 0x96)
    grey = RGBColor(0x55, 0x55, 0x55)

    document = Document()
    for section in document.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Name
    name_p = document.add_paragraph()
    name_run = name_p.add_run(candidate_name)
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = accent

    # Contact line with an accent rule underneath
    contact_p = document.add_paragraph()
    contact_run = contact_p.add_run(_cl_contact_text(links or [], location))
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = grey
    _add_bottom_border(contact_p, color="1A5C96")

    # Date
    date_p = document.add_paragraph(date.today().strftime("%B %d, %Y"))
    date_p.runs[0].font.size = Pt(11)

    # Recipient block
    who_p = document.add_paragraph()
    who_run = who_p.add_run("Hiring Manager")
    who_run.bold = True
    who_p.paragraph_format.space_after = Pt(0)
    document.add_paragraph(company_name)

    # Subject
    subject_p = document.add_paragraph()
    subject_run = subject_p.add_run(f"Re: Application for {role_title}")
    subject_run.bold = True
    subject_run.font.size = Pt(12)
    subject_run.font.color.rgb = accent

    # Salutation
    document.add_paragraph("Dear Hiring Manager,")

    # Body (justified)
    for para in body_paras:
        p = document.add_paragraph(para)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15

    # Sign-off
    document.add_paragraph("Warm regards,")
    sign_p = document.add_paragraph()
    sign_run = sign_p.add_run(candidate_name)
    sign_run.bold = True
    sign_run.font.color.rgb = accent
    sign_p.paragraph_format.space_after = Pt(0)
    tagline = headline or role_title
    if tagline:
        tag_p = document.add_paragraph()
        tag_run = tag_p.add_run(tagline)
        tag_run.font.size = Pt(10)
        tag_run.font.color.rgb = grey

    document.save(output_path)
    return output_path


def upload_document_to_gcs(local_path: str, gcs_path: str) -> str:
    settings = get_settings()
    if not settings.gcs_bucket_name:
        raise RuntimeError("GCS_BUCKET_NAME is not configured")

    bucket = get_gcs_client().bucket(settings.gcs_bucket_name)
    blob = bucket.blob(gcs_path)
    blob.upload_from_filename(local_path)

    return blob.generate_signed_url(
        version="v4",
        expiration=SIGNED_URL_EXPIRY,
        method="GET",
    )


def fresh_signed_url(gcs_path: str) -> str:
    """Issue a new signed URL for an existing object. Used to refresh expired download links."""
    settings = get_settings()
    if not settings.gcs_bucket_name:
        raise RuntimeError("GCS_BUCKET_NAME is not configured")
    bucket = get_gcs_client().bucket(settings.gcs_bucket_name)
    blob = bucket.blob(gcs_path)
    return blob.generate_signed_url(
        version="v4",
        expiration=SIGNED_URL_EXPIRY,
        method="GET",
    )


def delete_gcs_prefix(prefix: str) -> int:
    """Delete every object under a GCS prefix (e.g. ``pipelines/<id>/``).

    Best-effort and synchronous (run it via ``asyncio.to_thread`` from async
    endpoints). Returns the number of blobs deleted; a missing bucket or an
    individual delete failure is swallowed so cleanup never blocks the caller's
    primary action (deleting the Firestore record).
    """
    settings = get_settings()
    if not settings.gcs_bucket_name:
        return 0
    bucket = get_gcs_client().bucket(settings.gcs_bucket_name)
    deleted = 0
    for blob in bucket.list_blobs(prefix=prefix):
        try:
            blob.delete()
            deleted += 1
        except Exception:  # pragma: no cover - storage hiccup shouldn't block delete
            pass
    return deleted
